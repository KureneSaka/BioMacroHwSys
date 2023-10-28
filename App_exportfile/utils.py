from docx.table import Table, _Cell, _Row
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml, OxmlElement
from docx.shared import Inches, Length
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from App_dataSystem.models import *
from App_adminSystem.views.utils import quesList2dict

colorDict = {"table-primary": "C3DFFF",
             "table-secondary": "F9FBFF",
             "table-tertiary": "BFC4C9",
             "table-light": "F9FAFB",
             "table-info": "D8F9FB",
             "table-success": "B8EFD1",
             "table-danger": "FFB8C6",
             "thead-light": "E9ECEf",
             }


def setCellColor(cell: _Cell, colorStr: str):
    shading_list = locals()
    shading_list['shading_elm_'+str(cell)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    cell._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_'+str(cell)])


def setAlignment(cell: _Cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER


def setFont(cell: _Cell):
    cell.paragraphs[0].paragraph_format.space_after = Length(0)
    for r in cell.paragraphs[0].runs:
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

